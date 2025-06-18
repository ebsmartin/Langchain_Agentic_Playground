import sys
import os
from dotenv import load_dotenv

# Add the parent directory to the Python path so we can import from the main project
sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from conversation_parser import analyze_conversation_and_find_linkedin_profiles

def run_conversation_tests(test_number: int = None):
    """
    Run comprehensive conversation parser tests with different input types.
    
    Args:
        test_number (int, optional): Specific test number to run (1-6). 
                                   If None, runs all tests.
    """
    load_dotenv()
    
    # Define user identity with more comprehensive data
    user_identity = {
        'name': 'Eric Burton Martin',
        'title': 'Healthcare Data Analyst', 
        'company': 'Cognizant',
        'school': 'Colorado State University',
        'linkedin_url': 'https://www.linkedin.com/in/ebsmartin/'
    }
    
    # Test conversation text
    conversation_text = """
Person A: Hey, excuse me, is this seat taken? The keynote on DeFi scalability was packed!
Person B: Oh, hey! No, go for it. I'm Tanner, by the way. Yeah, that was a great talk.
Person A: I'm Eric. Nice to meet you, Tanner. So, what brings you to CryptoCon? Are you deep in the DeFi space too?
Person B: You could say that! I'm a Software Engineer at 1Inch. We're trying to make it easier for users to control their own data using blockchain. How about you, Jamie?
Person A: That's fascinating! I'm on the other side of the coin, so to speak. I'm an Associate Market Analyst at 1Inch. I spend my days tracking trends, tokenomics, and trying to predict the next big wave. So, less building, more observing.
Person B: Ah, a market guru! We definitely need more people like you to make sense of all this chaos. It's interesting, your analysis probably informs the kind of projects that get traction, which then impacts what developers like me build.
Person A: Exactly! It's all interconnected. I've been particularly interested in the intersection of NFTs and real-world asset tokenization lately. Seems like a huge potential area.
Person B: Totally agree. We've actually been exploring how decentralized identity could play a role in verifying ownership of tokenized assets. Makes the whole process more secure and transparent.
Person A: That's a great point. We should definitely connect. I'd love to hear more about your project and how DID can be applied there. Maybe we could grab a coffee tomorrow morning before the sessions start?
Person B: Sounds like a plan, Eric. I'm free around 8:30. Let's exchange contact info. Here's my LinkedIn QR code.
Person A: Perfect. Just scanned it. I'll send you a message. Looking forward to it!
"""
    
    print("="*60)
    print("CONVERSATION PARSER COMPREHENSIVE TESTING")
    print("="*60)
    print(f"USER IDENTITY: {user_identity['name']} - {user_identity['title']} at {user_identity['company']}")
    
    if test_number:
        print(f"RUNNING SINGLE TEST: {test_number}")
    else:
        print("RUNNING ALL TESTS")
    
    print("="*60)
    
    # Test 1: Direct Text Input
    if test_number is None or test_number == 1:
        print("\n🧪 TEST 1: DIRECT TEXT INPUT")
        print("-" * 40)
        
        result1 = analyze_conversation_and_find_linkedin_profiles(
            input_data=conversation_text,
            user_identity=user_identity,
            is_file_path=False,
            conversation_date='2024-06-18'
        )
        
        print(f"\n✅ Test 1 Results:")
        print(f"   - User excluded: {result1['user_excluded']}")
        print(f"   - Search queries found: {len(result1['search_queries'])}")
        print(f"   - LinkedIn profiles found: {sum(1 for p in result1['linkedin_profiles'] if p.get('linkedin_url') and 'Could not find' not in p['linkedin_url'])}")
    
    # Test 2: Text File Input
    if test_number is None or test_number == 2:
        print("\n🧪 TEST 2: TEXT FILE INPUT")
        print("-" * 40)
        
        # Create a test text file with clearer content
        test_text_file = "test_files/conversation.txt"
        os.makedirs("test_files", exist_ok=True)
        
        with open(test_text_file, 'w', encoding='utf-8') as f:
            f.write("""
Person A: Hi! I'm Matt, software engineer at Nickel5.
Person B: Nice to meet you Matt! I'm Eric Burton Martin, I do healthcare data analytics at Cognizant. I graduated from Colorado State University.
Person A: That's awesome! How long have you been at Cognizant?
Person B: About 5 months now. I love working on AI and big data projects for healthcare. What about you at Nickel5?
Person A: I've been there for 1 year, mainly working on improving revenue optimization algorithms. We should connect on LinkedIn!
""")
        
        print(f"📄 Test file content:")
        with open(test_text_file, 'r', encoding='utf-8') as f:
            print(f"   {f.read().strip()}")
        
        try:
            result2 = analyze_conversation_and_find_linkedin_profiles(
                input_data=test_text_file,
                user_identity=user_identity,
                is_file_path=True,
                conversation_date='2024-06-18'
            )
            
            print(f"\n✅ Test 2 Results:")
            print(f"   - User excluded: {result2['user_excluded']}")
            print(f"   - Search queries found: {len(result2['search_queries'])}")
            print(f"   - LinkedIn profiles found: {sum(1 for p in result2['linkedin_profiles'] if p.get('linkedin_url') and 'Could not find' not in p['linkedin_url'])}")
            
        except Exception as e:
            print(f"❌ Test 2 Failed: {e}")
    
    # Test 3: Audio File Input
    if test_number is None or test_number == 3:
        print("\n🧪 TEST 3: AUDIO FILE INPUT")
        print("-" * 40)
        
        audio_file = "test_files/conversation.mp3"
        if os.path.exists(audio_file):
            try:
                result3 = analyze_conversation_and_find_linkedin_profiles(
                    input_data=audio_file,
                    user_identity=user_identity,
                    is_file_path=True,
                    conversation_date='2024-06-18'
                )
                
                print(f"\n✅ Test 3 Results:")
                print(f"   - User excluded: {result3['user_excluded']}")
                print(f"   - Search queries found: {len(result3['search_queries'])}")
                print(f"   - LinkedIn profiles found: {sum(1 for p in result3['linkedin_profiles'] if p['linkedin_url'])}")
                
            except Exception as e:
                print(f"❌ Test 3 Failed: {e}")
        else:
            print("⏭️ Test 3 Skipped: No audio file found at test_files/conversation.mp3")
            print("   To test audio: Add an MP3 file with conversation to test_files/conversation.mp3")
    
    # Test 4: Image File Input
    if test_number is None or test_number == 4:
        print("\n🧪 TEST 4: IMAGE FILE INPUT")
        print("-" * 40)
        
        image_file = "test_files/business_card.png"
        if os.path.exists(image_file):
            try:
                result4 = analyze_conversation_and_find_linkedin_profiles(
                    input_data=image_file,
                    user_identity=user_identity,
                    is_file_path=True,
                    conversation_date='2024-06-18'
                )
                
                print(f"\n✅ Test 4 Results:")
                print(f"   - User excluded: {result4['user_excluded']}")
                print(f"   - Search queries found: {len(result4['search_queries'])}")
                print(f"   - LinkedIn profiles found: {sum(1 for p in result4['linkedin_profiles'] if p['linkedin_url'])}")
                
            except Exception as e:
                print(f"❌ Test 4 Failed: {e}")
        else:
            print("⏭️ Test 4 Skipped: No image file found at test_files/business_card.png")
            print("   To test image: Add a PNG/JPG file with conversation text to test_files/")
    
    # Test 5: PDF File Input  
    if test_number is None or test_number == 5:
        print("\n🧪 TEST 5: PDF FILE INPUT")
        print("-" * 40)
        
        pdf_file = "test_files/conversation.pdf"
        if os.path.exists(pdf_file):
            try:
                result5 = analyze_conversation_and_find_linkedin_profiles(
                    input_data=pdf_file,
                    user_identity=user_identity,
                    is_file_path=True,
                    conversation_date='2024-06-18'
                )
                
                print(f"\n✅ Test 5 Results:")
                print(f"   - User excluded: {result5['user_excluded']}")
                print(f"   - Search queries found: {len(result5['search_queries'])}")
                print(f"   - LinkedIn profiles found: {sum(1 for p in result5['linkedin_profiles'] if p['linkedin_url'])}")
                
            except Exception as e:
                print(f"❌ Test 5 Failed: {e}")
        else:
            print("⏭️ Test 5 Skipped: No PDF file found at test_files/conversation.pdf")
            print("   To test PDF: Add a PDF file with conversation text to test_files/")
    
    # Test 6: Word Document Input
    if test_number is None or test_number == 6:
        print("\n🧪 TEST 6: WORD DOCUMENT INPUT (.docx)")
        print("-" * 40)
        
        # Create a test Word document
        test_docx_file = "test_files/conversation.docx"
        
        try:
            # Create the Word document using python-docx
            from docx import Document
            
            doc = Document()
            doc.add_heading('Meeting Notes - Tech Conference', 0)
            doc.add_paragraph('Date: June 18, 2024')
            doc.add_paragraph('Location: CryptoCon 2024')
            doc.add_paragraph('')
            doc.add_heading('Conversation Log:', level=1)
            doc.add_paragraph(
                "Person A: Hi there! I'm Dr. Indrakshi Ray, Professor of cybersecurity at Colorado State University and Director of the Colorado Center for Cybersecurity. "
                "I specialize in all things cybersecurity."
            )
            doc.add_paragraph(
                "Person B: Nice to meet you Dr. Ray! I'm Eric Martin, I do AI consulting at Cognizant. "
                "I graduated from CSU and have been working on all sorts of AI integration projects."
            )
            doc.add_paragraph(
                "Person A: That's fascinating! How long have you been at Cognizant?"
            )
            doc.add_paragraph(
                "Person B: About 5 months now. I love working on cutting-edge technology products. "
                "What about your work at CSU?"
            )
            doc.add_paragraph(
                "Person A: I've been there for 23 years, mainly leading the cybersec research team. "
                "We should definitely connect and explore potential collaborations!"
            )
        
            doc.save(test_docx_file)
            print(f"📄 Created test Word document: {test_docx_file}")
            
            # Now test the analysis
            result6 = analyze_conversation_and_find_linkedin_profiles(
                input_data=test_docx_file,
                user_identity=user_identity,
                is_file_path=True,
                conversation_date='2025-06-18'
            )
            
            print(f"\n✅ Test 6 Results:")
            print(f"   - User excluded: {result6['user_excluded']}")
            print(f"   - Search queries found: {len(result6['search_queries'])}")
            print(f"   - LinkedIn profiles found: {sum(1 for p in result6['linkedin_profiles'] if p['linkedin_url'])}")
            print(f"   - Tool used: extract_text_from_word (via agent)")
            
            # Show what queries were generated
            if result6['search_queries']:
                print("   - Generated queries:")
                for query in result6['search_queries']:
                    print(f"     • {query}")
            
        except ImportError:
            print("❌ Test 6 Failed: python-docx not installed")
            print("   Install with: pip install python-docx")
        except Exception as e:
            print(f"❌ Test 6 Failed: {e}")
    
    # Only show summary if running all tests
    if test_number is None:
        print("\n" + "="*60)
        print("TESTING COMPLETED")
        print("="*60)
        
        # Summary of what each test demonstrates
        print("\n📋 TEST SUMMARY:")
        print("   Test 1: Direct text → Gemini native processing")
        print("   Test 2: .txt file → Simple file read + Gemini processing") 
        print("   Test 3: .mp3 audio → Gemini native audio processing")
        print("   Test 4: .png image → Gemini native vision processing")
        print("   Test 5: .pdf file → Gemini native PDF processing")
        print("   Test 6: .docx file → Agent + extract_text_from_word tool")
        print("\n🎯 Only Test 6 uses our custom tool and agent!")
    else:
        print(f"\n✅ Test {test_number} completed!")

if __name__ == "__main__":
    import sys
    
    # Check if a test number was provided as command line argument
    if len(sys.argv) > 1:
        try:
            test_num = int(sys.argv[1])
            if 1 <= test_num <= 6:
                print(f"Starting Conversation Parser Test {test_num}...")
                run_conversation_tests(test_num)
            else:
                print("❌ Error: Test number must be between 1 and 6")
                print("Usage: python test_conversation_parser.py [test_number]")
                print("Example: python test_conversation_parser.py 3")
        except ValueError:
            print("❌ Error: Please provide a valid test number (1-6)")
            print("Usage: python test_conversation_parser.py [test_number]")
    else:
        print("Starting Conversation Parser Testing...")
        run_conversation_tests()